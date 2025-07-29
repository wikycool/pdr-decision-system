#!/usr/bin/env python3
"""
Enhanced Document Analyzer for T&C Extraction
Specifically designed for enterprise PDR requirements
"""

import re
import json
import pandas as pd
from typing import Dict, List, Any, Tuple, Optional
from pathlib import Path
import logging
from datetime import datetime

# Optional imports for PDF processing
try:
    import pdfplumber
    import camelot
    import tabula
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    print("Warning: PDF processing libraries not available")

# Optional imports for DOCX processing
try:
    from docx import Document
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    print("Warning: DOCX processing library not available")

class EnhancedDocumentAnalyzer:
    """
    Enhanced document analyzer for T&C extraction with primary/secondary/tertiary categorization.
    Designed specifically for enterprise PDR requirements.
    """
    
    def __init__(self):
        self.setup_logging()
        self.setup_keywords()
    
    def setup_logging(self):
        """Setup logging configuration."""
        try:
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler('enhanced_analysis.log'),
                    logging.StreamHandler()
                ]
            )
            self.logger = logging.getLogger(__name__)
        except Exception as e:
            print(f"Warning: Could not setup logging: {e}")
            self.logger = None
    
    def setup_keywords(self):
        """Setup keyword categories for T&C extraction."""
        
        # Primary Keywords (High Priority - Must Have)
        self.primary_keywords = {
            'compliance': [
                'iso', 'astm', 'asme', 'ieee', 'ul', 'ce', 'fcc', 'rohs', 'reach',
                'certification', 'standard', 'requirement', 'specification',
                'regulation', 'guideline', 'protocol', 'framework', 'approval',
                'certified', 'qualified', 'tested', 'verified', 'validated',
                'approved', 'accepted', 'mandatory', 'obligatory', 'required'
            ],
            'technical': [
                'material', 'component', 'specification', 'dimension', 'tolerance',
                'performance', 'reliability', 'durability', 'safety', 'quality',
                'testing', 'validation', 'verification', 'manufacturing', 'assembly',
                'design', 'engineering', 'technical', 'mechanical', 'electrical',
                'structural', 'functional', 'operational', 'maintenance', 'service',
                'installation', 'configuration', 'setup', 'calibration', 'inspection'
            ],
            'cost': [
                'cost', 'price', 'budget', 'economic', 'financial', 'expense',
                'investment', 'roi', 'value', 'efficiency', 'optimization',
                'dollar', 'currency', 'payment', 'invoice', 'quote', 'estimate',
                'pricing', 'rate', 'fee', 'charge', 'expense', 'overhead'
            ]
        }
        
        # Secondary Keywords (Medium Priority - Should Have)
        self.secondary_keywords = {
            'quality': [
                'quality', 'assurance', 'control', 'inspection', 'testing',
                'verification', 'validation', 'certification', 'audit',
                'review', 'assessment', 'evaluation', 'check', 'monitor'
            ],
            'timeline': [
                'schedule', 'timeline', 'deadline', 'milestone', 'delivery',
                'duration', 'timeframe', 'period', 'phase', 'stage',
                'start', 'end', 'completion', 'target', 'due'
            ],
            'communication': [
                'communication', 'reporting', 'notification', 'update',
                'meeting', 'review', 'discussion', 'consultation',
                'feedback', 'approval', 'authorization', 'signature'
            ]
        }
        
        # Tertiary Keywords (Low Priority - Nice to Have)
        self.tertiary_keywords = {
            'documentation': [
                'documentation', 'manual', 'guide', 'instruction', 'procedure',
                'process', 'workflow', 'methodology', 'approach', 'strategy'
            ],
            'support': [
                'support', 'maintenance', 'service', 'help', 'assistance',
                'training', 'education', 'knowledge', 'expertise', 'consultation'
            ],
            'future': [
                'future', 'upgrade', 'enhancement', 'improvement', 'development',
                'evolution', 'advancement', 'innovation', 'modernization'
            ]
        }
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using multiple methods."""
        if not PDF_PROCESSING_AVAILABLE:
            print("Warning: PDF processing not available")
            return ""
        
        text = ""
        
        try:
            # Try pdfplumber first (best for text extraction)
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    print(f"Successfully extracted text using pdfplumber")
                    return text
            except Exception as e:
                print(f"pdfplumber failed: {e}")
            
            # Try camelot for tables
            try:
                tables = camelot.read_pdf(pdf_path, pages='all')
                for table in tables:
                    text += table.df.to_string() + "\n"
                print(f"Successfully extracted tables using camelot")
            except Exception as e:
                print(f"camelot failed: {e}")
            
            # Try tabula as fallback
            if not text.strip():
                try:
                    tables = tabula.read_pdf(pdf_path, pages='all')
                    for table in tables:
                        text += table.to_string() + "\n"
                    print(f"Successfully extracted tables using tabula")
                except Exception as e:
                    print(f"tabula failed: {e}")
            
            return text
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Error extracting text from PDF: {e}")
            else:
                print(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_PROCESSING_AVAILABLE:
            print("Warning: DOCX processing not available")
            return ""
        
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""

    def categorize_requirements(self, text: str) -> Dict[str, Dict[str, List[str]]]:
        """Categorize requirements into primary, secondary, and tertiary."""
        
        requirements = {
            'primary': {'compliance': [], 'technical': [], 'cost': []},
            'secondary': {'quality': [], 'timeline': [], 'communication': []},
            'tertiary': {'documentation': [], 'support': [], 'future': []}
        }
        
        lines = text.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            line_lower = line.lower()
            
            # Check primary categories
            for category, keywords in self.primary_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    requirements['primary'][category].append(line.strip())
                    break
            
            # Check secondary categories
            for category, keywords in self.secondary_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    requirements['secondary'][category].append(line.strip())
                    break
            
            # Check tertiary categories
            for category, keywords in self.tertiary_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    requirements['tertiary'][category].append(line.strip())
                    break
        
        return requirements
    
    def generate_sme_summary(self, requirements: Dict[str, Dict[str, List[str]]]) -> Dict[str, Any]:
        """Generate summary for SMEs and architects."""
        
        summary = {
            'total_requirements': 0,
            'primary_count': 0,
            'secondary_count': 0,
            'tertiary_count': 0,
            'critical_areas': [],
            'recommendations': [],
            'risk_assessment': {},
            'decision_points': []
        }
        
        # Count requirements by priority
        for priority, categories in requirements.items():
            for category, reqs in categories.items():
                summary['total_requirements'] += len(reqs)
                if priority == 'primary':
                    summary['primary_count'] += len(reqs)
                elif priority == 'secondary':
                    summary['secondary_count'] += len(reqs)
                elif priority == 'tertiary':
                    summary['tertiary_count'] += len(reqs)
        
        # Identify critical areas
        critical_categories = []
        for priority, categories in requirements.items():
            for category, reqs in categories.items():
                if len(reqs) > 0:
                    if priority == 'primary':
                        critical_categories.append(f"CRITICAL: {category.upper()} ({len(reqs)} requirements)")
                    elif priority == 'secondary':
                        critical_categories.append(f"IMPORTANT: {category.upper()} ({len(reqs)} requirements)")
                    else:
                        critical_categories.append(f"OPTIONAL: {category.upper()} ({len(reqs)} requirements)")
        
        summary['critical_areas'] = critical_categories
        
        # Generate recommendations
        if summary['primary_count'] > 0:
            summary['recommendations'].append("Focus on primary requirements first - these are mandatory")
        if summary['secondary_count'] > 0:
            summary['recommendations'].append("Address secondary requirements based on available resources")
        if summary['tertiary_count'] > 0:
            summary['recommendations'].append("Consider tertiary requirements for future enhancements")
        
        # Risk assessment
        summary['risk_assessment'] = {
            'high_risk': summary['primary_count'] == 0,
            'medium_risk': summary['secondary_count'] == 0,
            'low_risk': summary['tertiary_count'] == 0
        }
        
        # Decision points
        summary['decision_points'] = [
            "Which primary requirements are most critical?",
            "What resources are needed for secondary requirements?",
            "How to prioritize requirements across categories?",
            "What is the timeline for implementation?",
            "What are the cost implications?"
        ]
        
        return summary
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Analyze a document and extract T&C with categorization."""
        print(f"Starting enhanced analysis of: {file_path}")
        
        # Check if file exists
        if not Path(file_path).exists():
            print(f"File not found: {file_path}")
            return {}
        
        # Extract text
        text = ""
        try:
            if file_path.lower().endswith('.pdf'):
                text = self.extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_path.lower().endswith('.docx'):
                text = self.extract_text_from_docx(file_path)
            else:
                print(f"Unsupported file type: {file_path}")
                return {}
        except Exception as e:
            print(f"Error reading file: {e}")
            return {}
        
        if not text.strip():
            print("No text extracted from document")
            return {}
        
        # Categorize requirements
        requirements = self.categorize_requirements(text)
        
        # Generate SME summary
        summary = self.generate_sme_summary(requirements)
        
        # Compile results
        results = {
            'file_path': file_path,
            'text_length': len(text),
            'requirements': requirements,
            'sme_summary': summary,
            'analysis_timestamp': datetime.now().isoformat()
        }
        
        print(f"Enhanced analysis complete. Found {summary['total_requirements']} requirements")
        return results
    
    def analyze_multiple_documents(self, file_paths: List[str]) -> Dict[str, Any]:
        """Analyze multiple documents and combine results."""
        print(f"Analyzing {len(file_paths)} documents...")
        
        all_results = []
        combined_requirements = {
            'primary': {'compliance': [], 'technical': [], 'cost': []},
            'secondary': {'quality': [], 'timeline': [], 'communication': []},
            'tertiary': {'documentation': [], 'support': [], 'future': []}
        }
        
        for file_path in file_paths:
            result = self.analyze_document(file_path)
            if result:
                all_results.append(result)
                
                # Combine requirements
                for priority, categories in result['requirements'].items():
                    for category, reqs in categories.items():
                        combined_requirements[priority][category].extend(reqs)
        
        # Generate overall summary
        overall_summary = self.generate_sme_summary(combined_requirements)
        overall_summary['documents_analyzed'] = len(all_results)
        overall_summary['analysis_method'] = 'Enhanced T&C Extraction'
        
        return {
            'individual_results': all_results,
            'combined_requirements': combined_requirements,
            'overall_summary': overall_summary
        }

def main():
    """Test the enhanced document analyzer."""
    analyzer = EnhancedDocumentAnalyzer()
    
    # Test with sample text
    sample_text = """
    This is a test document with various T&C requirements:
    
    PRIMARY REQUIREMENTS:
    1. The component must meet ISO 9001 standards (COMPLIANCE)
    2. Material specifications should include ASTM A36 steel (TECHNICAL)
    3. Cost should not exceed $500 per unit (COST)
    
    SECONDARY REQUIREMENTS:
    4. Quality assurance procedures must be followed (QUALITY)
    5. Project timeline is 6 months (TIMELINE)
    6. Weekly status reports required (COMMUNICATION)
    
    TERTIARY REQUIREMENTS:
    7. User manual documentation (DOCUMENTATION)
    8. Training support provided (SUPPORT)
    9. Future upgrade path available (FUTURE)
    """
    
    print("Testing enhanced document analyzer...")
    requirements = analyzer.categorize_requirements(sample_text)
    summary = analyzer.generate_sme_summary(requirements)
    
    print(f"\nRequirements by Priority:")
    for priority, categories in requirements.items():
        print(f"\n{priority.upper()}:")
        for category, reqs in categories.items():
            print(f"  {category}: {len(reqs)} requirements")
    
    print(f"\nSME Summary:")
    print(f"Total Requirements: {summary['total_requirements']}")
    print(f"Critical Areas: {summary['critical_areas']}")
    print(f"Recommendations: {summary['recommendations']}")

if __name__ == "__main__":
    main() 