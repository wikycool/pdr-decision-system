#!/usr/bin/env python3
"""
LangChain-powered Document Analyzer for PDR System
"""

import os
import json
import re
import pandas as pd
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging

# LangChain imports
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    from langchain.prompts import PromptTemplate
    from langchain_openai import ChatOpenAI
    from langchain.chains import LLMChain
    from langchain.output_parsers import PydanticOutputParser
    from pydantic import BaseModel, Field
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    print(f"Warning: LangChain not available: {e}")
    LANGCHAIN_AVAILABLE = False

# Optional imports for DOCX processing
try:
    from docx import Document
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    print("Warning: DOCX processing library not available")

class Requirement(BaseModel):
    """Pydantic model for structured requirement extraction."""
    requirement_text: str = Field(description="The actual requirement text")
    category: str = Field(description="Category: compliance, technical, cost, or general")
    priority: str = Field(description="Priority: high, medium, or low")
    description: str = Field(description="Brief description of the requirement")

class RequirementsAnalysis(BaseModel):
    """Pydantic model for structured analysis output."""
    requirements: List[Requirement] = Field(description="List of extracted requirements")
    summary: Dict[str, Any] = Field(description="Summary statistics")
    key_insights: List[str] = Field(description="Key insights from the document")

class LangChainDocumentAnalyzer:
    """
    Enhanced document analyzer using LangChain for intelligent requirement extraction.
    """
    
    def __init__(self, openai_api_key: Optional[str] = None):
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        self.chain = None  # Initialize chain attribute
        self.setup_logging()
        
        if LANGCHAIN_AVAILABLE and self.openai_api_key:
            print("🤖 Initializing LangChain AI-powered analyzer...")
            self.setup_langchain()
            if self.chain:
                print("✅ LangChain AI analyzer ready!")
            else:
                print("⚠️ LangChain setup failed, using fallback mode")
        else:
            print("⚠️ LangChain or OpenAI API key not available. Using fallback mode.")
    
    def setup_logging(self):
        """Setup logging configuration."""
        try:
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler('langchain_analysis.log'),
                    logging.StreamHandler()
                ]
            )
            self.logger = logging.getLogger(__name__)
        except Exception as e:
            print(f"Warning: Could not setup logging: {e}")
            self.logger = None
    
    def setup_langchain(self):
        """Setup LangChain components."""
        try:
            print("🔧 Setting up LangChain components...")
            
            # Initialize LLM
            self.llm = ChatOpenAI(
                model_name="gpt-3.5-turbo",
                temperature=0.1,
                api_key=self.openai_api_key
            )
            print("✅ LLM initialized (GPT-3.5-turbo)")
            
            # Setup text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            print("✅ Text splitter configured")
            
            # Setup output parser
            self.output_parser = PydanticOutputParser(pydantic_object=RequirementsAnalysis)
            print("✅ Output parser configured")
            
            # Setup prompt template
            self.prompt_template = PromptTemplate(
                input_variables=["text"],
                template="""
                Analyze the following engineering document and extract requirements for a Preliminary Design Review (PDR).

                Document text:
                {text}

                Please extract all requirements and categorize them as:
                - compliance: standards, certifications, regulations (ISO, ASTM, ASME, etc.)
                - technical: specifications, materials, dimensions, performance
                - cost: budget, pricing, financial considerations
                - general: other important requirements

                For each requirement, provide:
                - requirement_text: the actual requirement
                - category: compliance/technical/cost/general
                - priority: high/medium/low
                - description: brief explanation

                Provide key insights about the document and summary statistics.

                {format_instructions}
                """
            )
            print("✅ Prompt template configured")

            # Setup chain
            self.chain = LLMChain(
                llm=self.llm,
                prompt=self.prompt_template,
                output_parser=self.output_parser
            )
            print("✅ LangChain chain created successfully")

        except Exception as e:
            print(f"❌ Error setting up LangChain: {e}")
            self.chain = None
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using multiple methods."""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            print(f"Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_PROCESSING_AVAILABLE:
            print("Warning: DOCX processing not available")
            return ""
        
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""

    def analyze_with_langchain(self, text: str) -> Dict[str, Any]:
        """Analyze text using LangChain."""
        if not self.chain:
            print("⚠️ LangChain not available, using fallback analysis")
            return self.fallback_analysis(text)

        try:
            print("🤖 Starting AI-powered analysis with LangChain...")

            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            print(f"📄 Split text into {len(chunks)} chunks for AI processing")

            all_requirements = []
            all_insights = []

            # Process each chunk
            for i, chunk in enumerate(chunks[:5]):  # Limit to first 5 chunks for cost
                print(f"🔄 Processing chunk {i+1}/{min(len(chunks), 5)} with AI...")

                try:
                    result = self.chain.run(text=chunk)

                    if hasattr(result, 'requirements'):
                        all_requirements.extend(result.requirements)
                        print(f"✅ Extracted {len(result.requirements)} requirements from chunk {i+1}")
                    if hasattr(result, 'key_insights'):
                        all_insights.extend(result.key_insights)

                except Exception as e:
                    print(f"⚠️ Error processing chunk {i+1}: {e}")
                    continue

            print(f"🎉 AI analysis complete! Found {len(all_requirements)} total requirements")

            # Compile results
            requirements_by_category = {
                'compliance': [],
                'technical': [],
                'cost': [],
                'general': []
            }

            for req in all_requirements:
                if hasattr(req, 'category') and hasattr(req, 'requirement_text'):
                    category = req.category.lower()
                    if category in requirements_by_category:
                        requirements_by_category[category].append(req.requirement_text)

            # Generate summary
            summary = {
                'total_requirements': sum(len(reqs) for reqs in requirements_by_category.values()),
                'compliance_count': len(requirements_by_category['compliance']),
                'technical_count': len(requirements_by_category['technical']),
                'cost_count': len(requirements_by_category['cost']),
                'general_count': len(requirements_by_category['general']),
                'key_insights': all_insights[:5],  # Limit insights
                'analysis_method': 'LangChain AI'
            }

            return {
                'requirements': requirements_by_category,
                'summary': summary,
                'raw_requirements': all_requirements
            }

        except Exception as e:
            print(f"❌ Error in LangChain analysis: {e}")
            return self.fallback_analysis(text)
    
    def fallback_analysis(self, text: str) -> Dict[str, Any]:
        """Fallback analysis when LangChain is not available."""
        print("🔍 Using keyword-based fallback analysis...")

        # Simple keyword-based analysis
        requirements = {
            'compliance': [],
            'technical': [],
            'cost': [],
            'general': []
        }

        lines = text.split('\n')

        compliance_keywords = ['iso', 'astm', 'asme', 'ieee', 'ul', 'ce', 'fcc', 'rohs', 'reach', 'certification', 'standard']
        technical_keywords = ['material', 'component', 'specification', 'dimension', 'tolerance', 'performance', 'reliability', 'safety', 'quality']
        cost_keywords = ['cost', 'price', 'budget', 'economic', 'financial', 'expense', 'dollar', 'currency']

        for line in lines:
            if not line.strip():
                continue

            line_lower = line.lower()

            # Categorize requirements
            if any(kw in line_lower for kw in compliance_keywords):
                requirements['compliance'].append(line.strip())
            elif any(kw in line_lower for kw in cost_keywords):
                requirements['cost'].append(line.strip())
            elif any(kw in line_lower for kw in technical_keywords):
                requirements['technical'].append(line.strip())
            elif re.search(r'^\d+\.|^[•\-*]|^[A-Z]\.', line.strip()):
                requirements['general'].append(line.strip())

        total_requirements = sum(len(reqs) for reqs in requirements.values())
        print(f"🔍 Keyword analysis found {total_requirements} requirements")

        summary = {
            'total_requirements': total_requirements,
            'compliance_count': len(requirements['compliance']),
            'technical_count': len(requirements['technical']),
            'cost_count': len(requirements['cost']),
            'general_count': len(requirements['general']),
            'key_insights': [f"Found {total_requirements} requirements using keyword analysis"],
            'analysis_method': 'Keyword-Based Fallback'
        }

        return {
            'requirements': requirements,
            'summary': summary
        }
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Analyze a document using LangChain."""
        print(f"📄 Starting analysis of: {file_path}")

        # Check if file exists
        if not Path(file_path).exists():
            print(f"❌ File not found: {file_path}")
            return {}

        # Extract text
        text = ""
        try:
            if file_path.lower().endswith('.pdf'):
                text = self.extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_path.lower().endswith('.docx'):
                text = self.extract_text_from_docx(file_path)
            else:
                print(f"❌ Unsupported file type: {file_path}")
                return {}
        except Exception as e:
            print(f"❌ Error reading file: {e}")
            return {}

        if not text.strip():
            print("❌ No text extracted from document")
            return {}

        print(f"📄 Extracted {len(text)} characters from document")

        # Analyze with LangChain
        results = self.analyze_with_langchain(text)

        # Add metadata
        results['file_path'] = file_path
        results['text_length'] = len(text)
        results['analysis_timestamp'] = pd.Timestamp.now().isoformat()

        analysis_method = results.get('summary', {}).get('analysis_method', 'Unknown')
        total_requirements = results.get('summary', {}).get('total_requirements', 0)
        
        print(f"✅ Analysis complete using {analysis_method}. Found {total_requirements} requirements")
        return results
    
    def analyze_multiple_documents(self, file_paths: List[str]) -> Dict[str, Any]:
        """Analyze multiple documents."""
        print(f"Analyzing {len(file_paths)} documents...")
        
        all_results = []
        combined_requirements = {
            'compliance': [],
            'technical': [],
            'cost': [],
            'general': []
        }
        
        for file_path in file_paths:
            result = self.analyze_document(file_path)
            if result:
                all_results.append(result)
                
                # Combine requirements
                for category in combined_requirements:
                    if category in result['requirements']:
                        combined_requirements[category].extend(result['requirements'][category])
        
        # Generate overall summary
        overall_summary = {
            'total_requirements': sum(len(reqs) for reqs in combined_requirements.values()),
            'compliance_count': len(combined_requirements['compliance']),
            'technical_count': len(combined_requirements['technical']),
            'cost_count': len(combined_requirements['cost']),
            'general_count': len(combined_requirements['general']),
            'key_insights': [
                f"Analyzed {len(all_results)} documents",
                f"Found {sum(len(reqs) for reqs in combined_requirements.values())} total requirements",
                f"Most common category: {max(combined_requirements.items(), key=lambda x: len(x[1]))[0]}"
            ]
        }
        
        return {
            'individual_results': all_results,
            'combined_requirements': combined_requirements,
            'overall_summary': overall_summary
        }

def main():
    """Test the LangChain analyzer."""
    analyzer = LangChainDocumentAnalyzer()
    
    # Test with sample text
    sample_text = """
    This is a test document with various requirements:
    
    1. The component must meet ISO 9001 standards
    2. Material specifications should include ASTM A36 steel
    3. Cost should not exceed $500 per unit
    4. Technical requirements include 0.1mm tolerance
    5. Performance testing is required
    6. Budget allocation is $10,000
    7. Safety certification needed
    8. Quality assurance procedures
    """
    
    print("Testing LangChain analyzer...")
    results = analyzer.analyze_with_langchain(sample_text)
    
    print(f"\nResults: {results}")

if __name__ == "__main__":
    main() 